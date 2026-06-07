"""Auto-generate operation report from recording input_log + screenshots.

Outputs three formats:
  - Markdown (.md)  — lightweight, git-friendly
  - HTML (.html)    — self-contained, interactive timeline
  - Word (.docx)    — formal document with embedded images
"""

from __future__ import annotations

import base64
import json
import os
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional

# Optional: PIL for image resize in Word
try:
    from PIL import Image
    _HAS_PIL = True
except ImportError:
    _HAS_PIL = False

# Optional: python-docx for Word output
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


# ══════════════════════════════════════════════════════════════════════
# 事件解析
# ══════════════════════════════════════════════════════════════════════

def parse_log(log_path: str) -> List[Dict]:
    """Parse JSONL input log into list of event dicts."""
    events = []
    with open(log_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            try:
                events.append(json.loads(line))
            except json.JSONDecodeError:
                pass
    return events


def classify_event(message: str) -> str:
    """Classify event message into a human-readable operation category."""
    if not message:
        return "unknown"
    # CONFIG / System
    if message.startswith("{"):
        return "system"
    # Mouse
    if "DblClick" in message:
        return "mouse"
    if "DragStart" in message or "DragMove" in message or "DragEnd" in message:
        return "drag"
    if "Click" in message or "Release" in message:
        return "mouse"
    if "ScrollUp" in message or "ScrollDown" in message:
        return "scroll"
    # Keyboard
    if "Hotkey" in message:
        return "hotkey"
    if "Key Press" in message or "Key Release" in message:
        return "keyboard"
    return "input"


def summarize_event(message: str) -> str:
    """Generate a concise Chinese description for an event."""
    if not message:
        return "未知操作"
    if message.startswith("{"):
        return "系统配置信息"
    # Drag
    if "DragStart" in message:
        return "拖拽开始"
    if "DragMove" in message:
        return "拖拽移动中"
    if "DragEnd" in message:
        return "拖拽结束"
    # Mouse
    if "DblClick" in message:
        return "双击"
    if "LClick" in message:
        return "左键单击"
    if "RClick" in message:
        return "右键单击"
    if "MClick" in message:
        return "中键单击"
    if "LRelease" in message:
        return "左键释放"
    if "RRelease" in message:
        return "右键释放"
    if "LRelease" in message:
        return "左键释放"
    # Scroll
    if "ScrollUp" in message:
        return "向上滚动"
    if "ScrollDown" in message:
        return "向下滚动"
    # Keyboard
    if "Hotkey" in message:
        parts = message.replace("Hotkey: ", "").strip()
        return f"组合键 [{parts}]"
    if "Key Press" in message:
        key = message.replace("Key Press: ", "").strip()
        return f"按键 [{key}]"
    if "Key Release" in message:
        key = message.replace("Key Release: ", "").strip()
        return f"释放按键 [{key}]"
    return message


def extract_coordinates(message: str) -> Optional[tuple]:
    """Extract (x, y) from message like 'LClick at (651, 527)'."""
    m = re.search(r"at \((\d+),\s*(\d+)\)", message)
    if m:
        return (int(m.group(1)), int(m.group(2)))
    return None


# ══════════════════════════════════════════════════════════════════════
# 事件分组：将连续相似事件（click/release、key press/release）合并
# ══════════════════════════════════════════════════════════════════════

def group_events(events: List[Dict]) -> List[Dict]:
    """Group consecutive click/release pairs into single user actions.
    
    Each group keeps the first screenshot, first timestamp, and merged message.
    """
    if not events:
        return []
    
    groups = []
    i = 0
    while i < len(events):
        ev = events[i]
        msg = ev.get("message", "")
        cat = classify_event(msg)
        
        # Skip system config
        if cat == "system":
            i += 1
            continue
        
        # Try to merge click+release into one action
        if cat == "mouse" and "Release" in msg:
            # This release might belong to the previous click
            if groups and classify_event(groups[-1].get("message", "")) == "mouse" and "Click" in groups[-1].get("message", ""):
                # Merge into previous group
                prev = groups[-1]
                prev["message"] = prev["message"].replace("Release", "").replace("Release", "")
                # Keep the release screenshot if the click didn't have one
                if not prev.get("screenshot") and ev.get("screenshot"):
                    prev["screenshot"] = ev["screenshot"]
                i += 1
                continue
        
        # Merge key press + release
        if cat == "keyboard" and "Key Release" in msg:
            if groups and classify_event(groups[-1].get("message", "")) == "keyboard" and "Key Press" in groups[-1].get("message", ""):
                # Keep press as the main event
                i += 1
                continue
        
        # Skip standalone releases (no preceding click in group)
        if cat == "mouse" and "Release" in msg:
            i += 1
            continue
        if cat == "keyboard" and "Key Release" in msg:
            i += 1
            continue
        
        groups.append(dict(ev))
        i += 1
    
    return groups


# ══════════════════════════════════════════════════════════════════════
# 辅助函数
# ══════════════════════════════════════════════════════════════════════

def img_to_base64(path: str, max_width: int = 1200) -> str:
    """Convert image to base64 data URI, resized for embedding."""
    if not path or not os.path.exists(path):
        return ""
    if _HAS_PIL:
        img = Image.open(path)
        w, h = img.size
        if w > max_width:
            ratio = max_width / w
            img = img.resize((max_width, int(h * ratio)), Image.LANCZOS if hasattr(Image, 'LANCZOS') else Image.ANTIALIAS)
        import io
        buf = io.BytesIO()
        fmt = path.lower().endswith('.png') and 'PNG' or 'JPEG'
        img.save(buf, format=fmt, quality=85)
        b64 = base64.b64encode(buf.getvalue()).decode()
        mime = "image/png" if fmt == "PNG" else "image/jpeg"
        return f"data:{mime};base64,{b64}"
    else:
        # Fallback: raw base64 without resize
        with open(path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        return f"data:image/png;base64,{b64}"


def get_screenshot_info(events: List[Dict], screenshot_dir: str) -> Dict:
    """Build screenshot metadata: which events reference which screenshots."""
    info = {}
    for idx, ev in enumerate(events):
        ss = ev.get("screenshot", "")
        if ss:
            info[ss] = {
                "event_index": idx,
                "timestamp": ev.get("timestamp", ""),
                "window": ev.get("window", ""),
                "message": ev.get("message", ""),
            }
    return info


def count_screenshots(screenshot_dir: str) -> int:
    """Count PNG files in screenshot directory."""
    if not os.path.isdir(screenshot_dir):
        return 0
    return len([f for f in os.listdir(screenshot_dir) if f.lower().endswith('.png')])


# ══════════════════════════════════════════════════════════════════════
# Markdown 报告
# ══════════════════════════════════════════════════════════════════════

def generate_markdown(events: List[Dict], screenshot_dir: str, output_path: str,
                     project_name: str = "", video_path: str = "") -> str:
    """Generate a Markdown report with timeline and embedded screenshots."""
    groups = group_events(events)
    ss_count = count_screenshots(screenshot_dir)
    
    lines = []
    lines.append(f"# 操作录屏报告 / Operation Recording Report")
    lines.append("")
    lines.append(f"**项目名称**: {project_name or 'N/A'}")
    lines.append(f"**操作步骤**: {len(groups)} 个用户动作")
    lines.append(f"**截图数量**: {ss_count}")
    lines.append("")
    
    # Timeline
    lines.append("---")
    lines.append("")
    lines.append("## 操作时间线 / Operation Timeline")
    lines.append("")
    
    used_screenshots = set()
    for idx, g in enumerate(groups, 1):
        ts = g.get("timestamp", "")
        win = g.get("window", "Unknown")
        msg = g.get("message", "")
        desc = summarize_event(msg)
        coords = extract_coordinates(msg)
        ss = g.get("screenshot", "")
        
        lines.append(f"### 步骤 {idx} — `[{ts}]`")
        lines.append("")
        lines.append(f"| 属性 | 值 |")
        lines.append(f"|------|------|")
        lines.append(f"| 操作 | **{desc}** |")
        lines.append(f"| 原始消息 | `{msg}` |")
        lines.append(f"| 活动窗口 | {win} |")
        if coords:
            lines.append(f"| 坐标 | ({coords[0]}, {coords[1]}) |")
        lines.append("")
        
        # Embed screenshot (only first time)
        if ss and ss not in used_screenshots:
            ss_full = os.path.join(os.path.dirname(output_path), ss)
            if os.path.exists(ss_full):
                lines.append(f"![步骤 {idx}]({ss})")
                lines.append("")
                used_screenshots.add(ss)
    
    lines.append("---")
    lines.append("")
    lines.append(f"*报告由 AgentRunner Recorder 自动生成*")
    
    md = "\n".join(lines)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(md)
    return md


# ══════════════════════════════════════════════════════════════════════
# HTML 报告（自包含，含交互时间线）
# ══════════════════════════════════════════════════════════════════════

_HTML_TEMPLATE = r"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
:root {{
  --bg: #0d1117;
  --surface: #161b22;
  --surface2: #1c2129;
  --border: #30363d;
  --text: #e6edf3;
  --text2: #8b949e;
  --accent: #58a6ff;
  --accent2: #7c8aff;
  --green: #3fb950;
  --red: #ff7b72;
  --orange: #f0883e;
  --mouse-color: #58a6ff;
  --key-color: #d2a8ff;
  --scroll-color: #3fb950;
  --drag-color: #f0883e;
  --hotkey-color: #ff7b72;
  --nav-w: 220px;
  --prop-w: 260px;
}}
[data-theme="light"] {{
  --bg: #f6f8fa;
  --surface: #ffffff;
  --surface2: #f0f2f5;
  --border: #d0d7de;
  --text: #1f2328;
  --text2: #656d76;
  --accent: #0969da;
  --accent2: #8250df;
  --green: #1a7f37;
  --red: #cf222e;
  --orange: #9a6700;
  --mouse-color: #0969da;
  --key-color: #8250df;
  --scroll-color: #1a7f37;
  --drag-color: #9a6700;
  --hotkey-color: #cf222e;
}}
.theme-toggle {{
  position: absolute; top: 12px; right: 12px;
  width: 28px; height: 28px;
  border: 1px solid var(--border);
  border-radius: 6px;
  background: var(--surface2);
  color: var(--text2);
  font-size: 14px; line-height: 1;
  cursor: pointer;
  display: flex; align-items: center; justify-content: center;
  transition: background 0.2s, color 0.2s, border-color 0.2s;
  z-index: 2;
}}
.theme-toggle:hover {{ background: var(--bg); color: var(--text); border-color: var(--accent); }}
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
html {{ scroll-behavior: smooth; }}
body {{
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
  background: var(--bg);
  color: var(--text);
  line-height: 1.6;
}}

/* ═════════════ Three-column layout ═════════════ */
.app {{ display: flex; min-height: 100vh; }}

/* ── Left nav sidebar ── */
.nav-sidebar {{
  position: fixed; top: 0; left: 0; bottom: 0;
  width: var(--nav-w);
  background: var(--surface);
  border-right: 1px solid var(--border);
  display: flex; flex-direction: column;
  z-index: 100;
}}
.nav-header {{
  padding: 16px 14px 12px;
  border-bottom: 1px solid var(--border);
  flex-shrink: 0;
}}
.nav-header h2 {{ font-size: 13px; font-weight: 700; }}
.nav-header .sub {{ font-size: 10px; color: var(--text2); margin-top: 2px; }}
.nav-stats {{
  display: flex; gap: 4px; flex-wrap: wrap;
  padding: 8px 14px;
  border-bottom: 1px solid var(--border);
  flex-shrink: 0;
}}
.nav-stats .chip {{
  display: flex; align-items: center; gap: 3px;
  font-size: 9px; color: var(--text2);
  background: var(--bg); border-radius: 10px; padding: 2px 6px;
}}
.nav-stats .dot {{ width: 5px; height: 5px; border-radius: 50%; }}
.nav-list {{
  flex: 1; overflow-y: auto; padding: 6px 0;
  scrollbar-width: thin; scrollbar-color: var(--border) transparent;
}}
.nav-list::-webkit-scrollbar {{ width: 3px; }}
.nav-list::-webkit-scrollbar-thumb {{ background: var(--border); border-radius: 2px; }}

.nav-item {{
  position: relative;
  padding: 8px 12px 8px 26px;
  cursor: pointer;
  transition: background 0.15s;
  border-left: 3px solid transparent;
}}
.nav-item:hover {{ background: var(--surface2); }}
.nav-item.active {{
  background: var(--surface2);
  border-left-color: var(--accent);
}}
.nav-item::before {{
  content: '';
  position: absolute; left: 17px; top: 0; bottom: 0;
  width: 1px; background: var(--border);
}}
.nav-item:first-child::before {{ top: 50%; }}
.nav-item:last-child::before {{ bottom: 50%; }}
.nav-item::after {{
  content: '';
  position: absolute; left: 13px; top: 50%;
  width: 9px; height: 9px; border-radius: 50%;
  border: 2px solid var(--border);
  background: var(--surface);
  transform: translateY(-50%);
  transition: all 0.2s;
}}
.nav-item.active::after {{
  border-color: var(--accent); background: var(--accent);
  box-shadow: 0 0 6px rgba(88,166,255,0.4);
}}
.nav-item.nav-mouse::after {{ border-color: var(--mouse-color); }}
.nav-item.nav-key::after {{ border-color: var(--key-color); }}
.nav-item.nav-scroll::after {{ border-color: var(--scroll-color); }}
.nav-item.nav-drag::after {{ border-color: var(--drag-color); }}
.nav-item.nav-hotkey::after {{ border-color: var(--hotkey-color); }}
.nav-item.active.nav-mouse::after {{ background: var(--mouse-color); }}
.nav-item.active.nav-key::after {{ background: var(--key-color); }}
.nav-item.active.nav-scroll::after {{ background: var(--scroll-color); }}
.nav-item.active.nav-drag::after {{ background: var(--drag-color); }}
.nav-item.active.nav-hotkey::after {{ background: var(--hotkey-color); }}
.nav-item.has-ss::after {{ width: 11px; height: 11px; }}

.nav-label {{ display: flex; align-items: center; gap: 6px; }}
.nav-num {{
  font-size: 9px; font-weight: 700; color: var(--accent);
  background: rgba(88,166,255,0.1); border-radius: 3px;
  padding: 1px 4px; flex-shrink: 0;
}}
.nav-action {{
  font-size: 11px; font-weight: 500; color: var(--text);
  white-space: nowrap; overflow: hidden; text-overflow: ellipsis;
}}
.nav-action.mouse {{ color: var(--mouse-color); }}
.nav-action.key {{ color: var(--key-color); }}
.nav-action.scroll {{ color: var(--scroll-color); }}
.nav-action.drag {{ color: var(--drag-color); }}
.nav-action.hotkey {{ color: var(--hotkey-color); }}
.nav-meta {{
  display: flex; align-items: center; gap: 6px; margin-top: 1px;
}}
.nav-time {{
  font-size: 9px; color: var(--text2);
  font-family: "SF Mono","Cascadia Code","Consolas",monospace;
}}
.nav-ss {{ font-size: 8px; color: var(--text2); background: var(--bg); border-radius: 3px; padding: 0 3px; }}

/* ── Center main area ── */
.main-area {{
  flex: 1;
  margin-left: var(--nav-w);
  margin-right: var(--prop-w);
  padding: 8px 12px 72px;
  max-width: calc(100% - var(--nav-w) - var(--prop-w));
  display: flex;
  flex-direction: column;
  overflow: hidden;
}}

.main-header {{
  text-align: center;
  margin-bottom: 10px;
  padding-bottom: 8px;
  border-bottom: 1px solid var(--border);
}}
.main-header h1 {{
  font-size: 16px; font-weight: 700; margin-bottom: 2px;
  background: linear-gradient(135deg, var(--accent), var(--accent2));
  -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}}
.main-header .meta {{
  display: flex; justify-content: center; gap: 8px; flex-wrap: wrap; margin-top: 4px;
}}
.main-header .meta-item {{
  background: var(--surface); border: 1px solid var(--border);
  border-radius: 6px; padding: 3px 10px; font-size: 11px;
}}
.main-header .meta-item .label {{ color: var(--text2); margin-right: 4px; }}
.main-header .meta-item .value {{ font-weight: 600; color: var(--accent); }}

/* ── Step cards (center) ── */
.step-list {{
  display: flex;
  flex-direction: column;
  flex: 1;
  align-items: center;
  justify-content: center;
}}
.step-card {{
  position: relative;
  border-radius: 12px;
  overflow: hidden;
  border: 1px solid var(--border);
  background: var(--surface);
  flex-shrink: 0;
  width: 100%;
  transition: border-color 0.3s, box-shadow 0.3s, opacity 0.35s ease, transform 0.35s ease;
}}
.step-card.page-active {{
  opacity: 1; transform: translateY(0);
  border-color: var(--accent);
  box-shadow: 0 0 0 1px var(--accent), 0 4px 24px rgba(88,166,255,0.1);
}}
.step-card.page-hidden {{
  opacity: 0; transform: translateY(30px);
  pointer-events: none; position: absolute; visibility: hidden;
}}
.step-card-top {{
  display: flex; align-items: center; gap: 10px;
  padding: 10px 16px;
  border-bottom: 1px solid var(--border);
  background: var(--surface2);
}}
.step-badge {{
  font-size: 10px; font-weight: 700; color: var(--bg);
  background: var(--accent); border-radius: 4px; padding: 2px 7px;
  flex-shrink: 0;
}}
.step-msg {{
  font-size: 13px; font-weight: 500; color: var(--text);
  flex: 1; white-space: nowrap; overflow: hidden; text-overflow: ellipsis;
}}
.step-msg.mouse {{ color: var(--mouse-color); }}
.step-msg.key {{ color: var(--key-color); }}
.step-msg.scroll {{ color: var(--scroll-color); }}
.step-msg.drag {{ color: var(--drag-color); }}
.step-msg.hotkey {{ color: var(--hotkey-color); }}
.step-ts {{
  font-size: 10px; color: var(--text2);
  font-family: "SF Mono","Cascadia Code","Consolas",monospace;
  flex-shrink: 0;
}}
.step-screenshot {{
  border-top: 1px solid var(--border);
}}
.step-screenshot img {{
  width: 100%; height: auto; display: block;
  margin: 0 auto;
  cursor: zoom-in; transition: max-width 0.2s, max-height 0.2s;
  max-height: calc(100vh - 160px);
  object-fit: contain;
  background: var(--bg);
}}
.step-screenshot img:hover {{ transform: scale(1.01); }}

/* ── Right property sidebar ── */
.prop-sidebar {{
  position: fixed; top: 0; right: 0; bottom: 0;
  width: var(--prop-w);
  background: var(--surface);
  border-left: 1px solid var(--border);
  display: flex; flex-direction: column;
  z-index: 100;
  overflow: hidden;
  transition: opacity 0.2s;
}}
.prop-header {{
  padding: 16px 16px 12px;
  border-bottom: 1px solid var(--border);
  flex-shrink: 0;
}}
.prop-header h3 {{
  font-size: 13px; font-weight: 700;
}}
.prop-header .prop-step-id {{
  font-size: 10px; color: var(--accent); margin-top: 2px;
  font-family: "SF Mono","Cascadia Code","Consolas",monospace;
}}
.prop-body {{
  flex: 1; overflow-y: auto; padding: 0;
  scrollbar-width: thin; scrollbar-color: var(--border) transparent;
}}
.prop-body::-webkit-scrollbar {{ width: 3px; }}
.prop-body::-webkit-scrollbar-thumb {{ background: var(--border); border-radius: 2px; }}

.prop-section {{
  padding: 14px 16px;
  border-bottom: 1px solid var(--border);
}}
.prop-section-title {{
  font-size: 10px; font-weight: 700; color: var(--text2);
  text-transform: uppercase; letter-spacing: 0.5px;
  margin-bottom: 10px;
}}
.prop-row {{
  display: flex;
  justify-content: space-between;
  align-items: flex-start;
  gap: 8px;
  margin-bottom: 8px;
}}
.prop-row:last-child {{ margin-bottom: 0; }}
.prop-key {{
  font-size: 11px; color: var(--text2); font-weight: 500;
  flex-shrink: 0; min-width: 56px;
}}
.prop-val {{
  font-size: 11px; color: var(--text);
  text-align: right; word-break: break-all;
}}
.prop-val code {{
  background: var(--bg); padding: 1px 5px; border-radius: 3px;
  font-size: 10px; color: var(--accent);
}}
.prop-action-tag {{
  display: inline-block;
  font-size: 11px; font-weight: 600;
  padding: 3px 10px; border-radius: 4px;
  margin-bottom: 10px;
}}
.prop-action-tag.mouse {{ background: rgba(88,166,255,0.12); color: var(--mouse-color); }}
.prop-action-tag.key {{ background: rgba(210,168,255,0.12); color: var(--key-color); }}
.prop-action-tag.scroll {{ background: rgba(63,185,80,0.12); color: var(--scroll-color); }}
.prop-action-tag.drag {{ background: rgba(240,136,62,0.12); color: var(--drag-color); }}
.prop-action-tag.hotkey {{ background: rgba(255,123,114,0.12); color: var(--hotkey-color); }}

/* ── Bottom footer with page navigation ── */
.bottom-footer {{
  position: fixed;
  bottom: 0;
  left: var(--nav-w);
  right: var(--prop-w);
  height: 56px;
  background: var(--surface);
  border-top: 1px solid var(--border);
  display: flex;
  align-items: center;
  justify-content: center;
  gap: 6px;
  z-index: 200;
  padding: 0 16px;
}}
.page-nav {{
  display: flex;
  gap: 6px;
  align-items: center;
}}
.page-nav-btn {{
  width: 36px; height: 36px;
  border: 1px solid var(--border);
  border-radius: 7px;
  background: var(--surface2);
  color: var(--text2);
  font-size: 16px; line-height: 1;
  cursor: pointer;
  display: flex; align-items: center; justify-content: center;
  transition: background 0.15s, color 0.15s, border-color 0.15s;
}}
.page-nav-btn:hover {{
  background: var(--accent);
  color: var(--bg);
  border-color: var(--accent);
}}
.page-nav-btn:disabled {{
  opacity: 0.3; cursor: default;
  background: var(--surface2); color: var(--text2); border-color: var(--border);
}}
.page-nav-sep {{
  width: 1px;
  background: var(--border);
  margin: 4px 2px;
}}
.page-nav-info {{
  display: flex; align-items: center;
  font-size: 11px; color: var(--text2);
  padding: 0 6px;
  font-family: "SF Mono","Cascadia Code","Consolas",monospace;
  white-space: nowrap;
}}
@media (max-width: 768px) {{
  .bottom-footer {{ left: 0; right: 0; height: 50px; }}
  .page-nav-btn {{ width: 32px; height: 32px; font-size: 14px; }}
}}

/* ── Lightbox ── */
.lightbox {{
  display: none; position: fixed;
  top: 0; left: 0; right: 0; bottom: 0;
  background: rgba(0,0,0,0.92);
  z-index: 1000; cursor: zoom-out;
  justify-content: center; align-items: center;
}}
.lightbox.active {{ display: flex; }}
.lightbox img {{
  max-width: 95vw; max-height: 95vh;
  object-fit: contain; border-radius: 8px;
}}

/* ── Footer (hidden, nav in bottom-footer) ── */
.main-footer {{ display: none; }}


/* ── Zoom slider ── */
.zoom-control {{
  display: flex; align-items: center; gap: 6px; margin-left: 12px;
  padding-left: 12px; border-left: 1px solid var(--border);
}}
.zoom-slider {{
  -webkit-appearance: none; appearance: none;
  width: 100px; height: 4px; border-radius: 2px;
  background: var(--border); outline: none; cursor: pointer;
}}
.zoom-slider::-webkit-slider-thumb {{
  -webkit-appearance: none; appearance: none;
  width: 14px; height: 14px; border-radius: 50%;
  background: var(--accent); cursor: pointer;
  border: 2px solid var(--bg);
  box-shadow: 0 1px 4px rgba(0,0,0,0.3);
}}
.zoom-slider::-moz-range-thumb {{
  width: 14px; height: 14px; border-radius: 50%;
  background: var(--accent); cursor: pointer;
  border: 2px solid var(--bg);
}}
.zoom-pct {{
  font-size: 11px; color: var(--text2); min-width: 32px; text-align: center;
}}
/* ── Responsive ── */
@media (max-width: 1100px) {{
  .prop-sidebar {{ display: none; }}
  .main-area {{ margin-right: 0; max-width: calc(100% - var(--nav-w)); }}
  .bottom-footer {{ right: 0; }}
}}
@media (max-width: 768px) {{
  .nav-sidebar {{ display: none; }}
  .main-area {{ margin-left: 0; max-width: 100%; padding: 12px 12px 70px; }}
  .bottom-footer {{ left: 0; right: 0; }}
}}
</style>
</head>
<body>
<div class="app">

<!-- Left nav sidebar -->
<aside class="nav-sidebar">
  <div class="nav-header" style="position:relative;">
    <h2>时间轴导航</h2>
    <button class="theme-toggle" id="themeToggle" title="切换主题">☀</button>
    <div class="sub">{total_steps} 步骤 / {ss_count} 截图</div>
  </div>
  <div class="nav-stats">{sidebar_stats}</div>
  <nav class="nav-list" id="navList">{nav_items_html}</nav>
</aside>

<!-- Center main area -->
<div class="main-area">
  <div class="main-header">
    <h1>{title}</h1>
    <div class="meta">
      <div class="meta-item"><span class="label">项目</span><span class="value">{project_name}</span></div>
      <div class="meta-item"><span class="label">步骤</span><span class="value">{total_steps}</span></div>
      <div class="meta-item"><span class="label">截图</span><span class="value">{ss_count}</span></div>
    </div>
  </div>

  <div class="step-list" id="stepList">
{steps_html}
  </div>

  <div class="main-footer">
    <p>报告由 <strong>AgentRunner Recorder</strong> 自动生成</p>
  </div>
</div>

<!-- Right property sidebar -->
<aside class="prop-sidebar" id="propSidebar">
  <div class="prop-header">
    <h3>步骤属性</h3>
    <div class="prop-step-id" id="propStepId">--</div>
  </div>
  <div class="prop-body" id="propBody">
    <div class="prop-section" style="text-align:center;color:var(--text2);font-size:14px;padding:40px 16px;">
      点击左侧步骤查看属性
    </div>
  </div>
</aside>

</div><!-- /.app -->

<!-- Bottom footer with page navigation -->
<div class="bottom-footer" id="bottomFooter">
  <div class="page-nav visible" id="pageNav">
    <button class="page-nav-btn" id="btnPrev" title="上一步 (↑)">▲</button>
    <button class="page-nav-btn" id="btnNext" title="下一步 (↓)">▼</button>
    <div class="page-nav-sep"></div>
    <button class="page-nav-btn" id="btnPrevPage" title="上一页 (←)">◀</button>
    <button class="page-nav-btn" id="btnNextPage" title="下一页 (→)">▶</button>
    <div class="page-nav-sep"></div>
    <div class="page-nav-info" id="pageInfo">1 / {total_steps}</div>
  </div>
  <div class="zoom-control">
    <span style="font-size:14px;color:var(--text2);">缩放</span>
    <input type="range" class="zoom-slider" id="zoomSlider" min="50" max="100" value="100" step="5" title="调整截图显示比例（控制四周留白）">
    <span class="zoom-pct" id="zoomPct">100%</span>
  </div>
  <div style="font-size:13px;color:var(--text2);margin-left:10px;white-space:nowrap;">AgentRunner Recorder</div>
</div>

<div class="lightbox" id="lightbox">
  <img src="" id="lightbox-img" alt="fullscreen">
</div>

<script>
// ── Data for right sidebar ──
const stepsData = {steps_data_json};

// ── Nav click → scroll + update properties ──
// ── Nav click handled by page-nav module ──

function showProperties(idx) {{
  if (idx < 0 || idx >= stepsData.length) return;
  const d = stepsData[idx];
  document.getElementById('propStepId').textContent = 'Step #' + d.num + '  ' + d.timestamp;
  let html = '';
  // Action
  html += '<div class=\"prop-section\"><div class=\"prop-action-tag ' + d.cat_class + '\">' + d.desc + '</div>';
  html += buildRows([['时间戳', d.timestamp], ['原始消息', '<code>' + escHtml(d.message) + '</code>']]);
  html += '</div>';
  // Window & coords
  html += '<div class=\"prop-section\"><div class=\"prop-section-title\">位置信息</div>';
  let locRows = [['活动窗口', escHtml(d.window)]];
  if (d.coords) locRows.push(['坐标', '(' + d.coords[0] + ', ' + d.coords[1] + ')']);
  html += buildRows(locRows) + '</div>';
  // Screenshot
  if (d.has_screenshot) {{
    html += '<div class=\"prop-section\"><div class=\"prop-section-title\">截图</div>';
    html += buildRows([['状态', '已捕获']]) + '</div>';
  }}
  document.getElementById('propBody').innerHTML = html;
}}
function buildRows(rows) {{
  return rows.map(r => '<div class=\"prop-row\"><span class=\"prop-key\">' + r[0] + '</span><span class=\"prop-val\">' + r[1] + '</span></div>').join('');
}}
function escHtml(s) {{
  return s.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
}}

// ── Scroll spy replaced by page-nav ──

// ── Theme toggle ──
(function() {{
  const btn = document.getElementById('themeToggle');
  const icons = {{ dark: '☀', light: '☾' }};
  function setTheme(t) {{
    document.documentElement.setAttribute('data-theme', t);
    btn.textContent = icons[t];
    localStorage.setItem('report-theme', t);
  }}
  const saved = localStorage.getItem('report-theme') ||
    (matchMedia('(prefers-color-scheme: light)').matches ? 'light' : 'dark');
  setTheme(saved);
  btn.addEventListener('click', () => {{
    const cur = document.documentElement.getAttribute('data-theme');
    setTheme(cur === 'dark' ? 'light' : 'dark');
  }});
  matchMedia('(prefers-color-scheme: light)').addEventListener('change', e => {{
    if (!localStorage.getItem('report-theme')) setTheme(e.matches ? 'light' : 'dark');
  }});
}})();

// ── Page navigation: one-step-per-view ──
(function() {{
  const stepCards = document.querySelectorAll('.step-card[id]');
  const total = stepCards.length;
  if (!total) return;
  const navItems = document.querySelectorAll('.nav-item');
  const pageNav = document.getElementById('pageNav');
  const pageInfo = document.getElementById('pageInfo');
  const btnPrev = document.getElementById('btnPrev');
  const btnNext = document.getElementById('btnNext');
  const btnPrevPage = document.getElementById('btnPrevPage');
  const btnNextPage = document.getElementById('btnNextPage');
  let currentStep = 0;

  function updatePageInfo(idx) {{
    pageInfo.textContent = (idx + 1) + ' / ' + total;
    btnPrev.disabled = idx <= 0;
    btnNext.disabled = idx >= total - 1;
  }}

  function showStep(idx) {{
    if (idx < 0 || idx >= total) return;
    currentStep = idx;
    stepCards.forEach((s, i) => {{
      if (i === idx) {{
        s.classList.remove('page-hidden');
        s.classList.add('page-active');
      }} else {{
        s.classList.remove('page-active');
        s.classList.add('page-hidden');
      }}
    }});
    window.scrollTo({{ top: 0, behavior: 'instant' }});
    updatePageInfo(idx);
    navItems.forEach(n => n.classList.remove('active'));
    if (navItems[idx]) {{
      navItems[idx].classList.add('active');
      const navList = document.getElementById('navList');
      const lr = navList.getBoundingClientRect();
      const er = navItems[idx].getBoundingClientRect();
      if (er.top < lr.top || er.bottom > lr.bottom)
        navItems[idx].scrollIntoView({{ behavior: 'smooth', block: 'nearest' }});
    }}
    if (typeof showProperties === 'function') showProperties(idx);
  }}

  // Navigation already visible via HTML class, overflow handled by CSS

  btnPrev.addEventListener('click', () => showStep(currentStep - 1));
  btnNext.addEventListener('click', () => showStep(currentStep + 1));
  const PAGE_JUMP = 3;
  btnPrevPage.addEventListener('click', () => showStep(Math.max(0, currentStep - PAGE_JUMP)));
  btnNextPage.addEventListener('click', () => showStep(Math.min(total - 1, currentStep + PAGE_JUMP)));

  document.addEventListener('keydown', (e) => {{
    if (e.target.tagName === 'INPUT' || e.target.tagName === 'TEXTAREA') return;
    if (e.key === 'ArrowUp' || e.key === 'k') {{ e.preventDefault(); showStep(currentStep - 1); }}
    else if (e.key === 'ArrowDown' || e.key === 'j') {{ e.preventDefault(); showStep(currentStep + 1); }}
    else if (e.key === 'ArrowLeft' || e.key === 'h') {{ e.preventDefault(); showStep(Math.max(0, currentStep - PAGE_JUMP)); }}
    else if (e.key === 'ArrowRight' || e.key === 'l') {{ e.preventDefault(); showStep(Math.min(total - 1, currentStep + PAGE_JUMP)); }}
    else if (e.key === 'Home') {{ e.preventDefault(); showStep(0); }}
    else if (e.key === 'End') {{ e.preventDefault(); showStep(total - 1); }}
  }});

  navItems.forEach((item, i) => {{
    item.addEventListener('click', (e) => {{ e.stopPropagation(); showStep(i); }});
  }});

  window._showStep = showStep;
  showStep(0);
}})();

// ── Lightbox ──
document.querySelectorAll('.step-screenshot img').forEach(img => {{
  img.addEventListener('click', () => {{
    const lb = document.getElementById('lightbox');
    document.getElementById('lightbox-img').src = img.dataset.full;
    lb.classList.add('active');
  }});
}});
document.getElementById('lightbox').addEventListener('click', () => {{
  document.getElementById('lightbox').classList.remove('active');
}});
document.addEventListener('keydown', e => {{
  if (e.key === 'Escape') document.getElementById('lightbox').classList.remove('active');
}});

// ── Zoom slider: control screenshot whitespace ──
(function() {{
  var slider = document.getElementById('zoomSlider');
  var pctLabel = document.getElementById('zoomPct');
  if (!slider) return;

  function applyZoom(val) {{
    document.querySelectorAll('.step-screenshot img').forEach(function(img) {{
      img.style.maxWidth = val + '%';
    }});
    pctLabel.textContent = val + '%';
    localStorage.setItem('report-zoom', val);
  }}

  var saved = localStorage.getItem('report-zoom');
  if (saved) {{
    slider.value = saved;
    applyZoom(parseInt(saved));
  }}

  slider.addEventListener('input', function() {{
    applyZoom(parseInt(this.value));
  }});
}})();

</script>
</body>
</html>"""
def generate_json(events, screenshot_dir, output_path, project_name="", video_path=""):
    """Generate a JSON report from grouped events."""
    groups = group_events(events)
    steps = []
    used_screenshots = set()
    for idx, g in enumerate(groups, 1):
        step = {
            "step": idx,
            "category": classify_event(g.get("message", "")),
            "description": summarize_event(g.get("message", "")),
            "timestamp": g.get("timestamp", ""),
            "message": g.get("message", ""),
            "window": g.get("window", ""),
        }
        coords = extract_coordinates(g.get("message", ""))
        if coords:
            step["coordinates"] = {"x": coords[0], "y": coords[1]}
        ss = g.get("screenshot", "")
        if ss and ss not in used_screenshots:
            ss_path = os.path.join(screenshot_dir, ss) if not os.path.isabs(ss) else ss
            if os.path.exists(ss_path):
                with open(ss_path, "rb") as f:
                    step["screenshot_base64"] = base64.b64encode(f.read()).decode()
                step["screenshot_file"] = ss
            else:
                step["screenshot_file"] = ss
            used_screenshots.add(ss)
        steps.append(step)

    report = {
        "project": project_name,
        "generated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total_steps": len(steps),
        "video": video_path or "",
        "steps": steps,
    }

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)


def generate_html(events: List[Dict], screenshot_dir: str, output_path: str,
                 project_name: str = "", video_path: str = "") -> str:
    """Generate a self-contained HTML report with 3-column layout:
    left nav, center screenshots+messages, right properties."""
    import html as html_mod

    groups = group_events(events)
    ss_count = count_screenshots(screenshot_dir)

    cat_counts = {}
    cat_labels = {"mouse": "鼠标", "keyboard": "键盘", "scroll": "滚动", "drag": "拖拽", "hotkey": "热键"}
    cat_dot_class = {"mouse": "mouse", "keyboard": "key", "scroll": "scroll", "drag": "drag", "hotkey": "hotkey"}
    for g in groups:
        cat = classify_event(g.get("message", ""))
        cat_counts[cat] = cat_counts.get(cat, 0) + 1

    # Sidebar stats
    sidebar_chips = []
    for cat, label in cat_labels.items():
        if cat in cat_counts:
            cls = cat.replace("keyboard", "key")
            sidebar_chips.append(
                f'<span class="chip"><span class="dot" style="background:var(--{cls}-color)"></span>{label} {cat_counts[cat]}</span>'
            )
    sidebar_stats = "".join(sidebar_chips)

    # Build steps (center) + nav items (left) + JS data (right)
    steps_html = []
    nav_items = []
    steps_data = []
    used_screenshots = set()
    report_base = os.path.dirname(output_path)

    for idx, g in enumerate(groups, 1):
        ts = g.get("timestamp", "")
        win = g.get("window", "")
        msg = g.get("message", "")
        desc = summarize_event(msg)
        cat = classify_event(msg)
        cat_class = cat_dot_class.get(cat, "mouse")
        coords = extract_coordinates(msg)
        ss = g.get("screenshot", "")
        has_ss = bool(ss and ss not in used_screenshots)

        # Center: step card
        ss_html = ""
        if ss and ss not in used_screenshots:
            ss_full = os.path.join(report_base, ss)
            if os.path.exists(ss_full):
                b64_uri = img_to_base64(ss_full, max_width=1200)
                if b64_uri:
                    ss_html = f"""<div class="step-screenshot">
  <img src="{b64_uri}" data-full="{b64_uri}" alt="Step {idx}" loading="lazy">
</div>"""
                    used_screenshots.add(ss)

        step = f"""<div class="step-card" id="step-{idx}">
  <div class="step-card-top">
    <span class="step-badge">#{idx}</span>
    <span class="step-msg {cat_class}">{html_mod.escape(desc)}</span>
    <span class="step-ts">{ts}</span>
  </div>
  {ss_html}
</div>"""
        steps_html.append(step)

        # Nav item
        ss_badge = '<span class="nav-ss">SS</span>' if has_ss else ""
        nav_items.append(
            f"""    <div class="nav-item nav-{cat_class}{" has-ss" if has_ss else ""}" data-target="step-{idx}">
      <div class="nav-label">
        <span class="nav-num">#{idx}</span>
        <span class="nav-action {cat_class}">{html_mod.escape(desc)}</span>
      </div>
      <div class="nav-meta">
        <span class="nav-time">{ts}</span>
        {ss_badge}
      </div>
    </div>"""
        )

        # JS data for right sidebar
        coords_json = f"[{coords[0]},{coords[1]}]" if coords else "null"
        steps_data.append(
            f'{{"num":{idx},"timestamp":"{ts}","message":"{html_mod.escape(msg).replace(chr(34), chr(92)+chr(34))}",'
            f'"window":"{html_mod.escape(win[:60])}","desc":"{html_mod.escape(desc)}",'
            f'"cat_class":"{cat_class}","coords":{coords_json},"has_screenshot":{str(has_ss).lower()}}}'
        )

    title = f"操作录屏报告 — {project_name}" if project_name else "操作录屏报告"
    html_out = _HTML_TEMPLATE.format(
        title=title,
        project_name=project_name or "N/A",
        total_steps=len(groups),
        ss_count=ss_count,
        sidebar_stats=sidebar_stats,
        nav_items_html="\n".join(nav_items),
        steps_html="\n".join(steps_html),
        steps_data_json="[" + ",".join(steps_data) + "]",
    )

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_out)
    return html_out

# ══════════════════════════════════════════════════════════════════════
# Word 报告
# ══════════════════════════════════════════════════════════════════════

def _set_run_font(run, font_name='微软雅黑', font_size=None, bold=False,
                  color=None):
    """Safely set run font properties including east asian font."""
    if font_size:
        run.font.size = font_size
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    # East asian font
    rpr = run.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)


def generate_word(events: List[Dict], screenshot_dir: str, output_path: str,
                  project_name: str = "", video_path: str = "") -> Optional[str]:
    """Generate a Word document report with embedded screenshots."""
    if not _HAS_DOCX:
        print("[report_generator] python-docx not installed, skipping Word report")
        return None
    
    groups = group_events(events)
    ss_count = count_screenshots(screenshot_dir)
    
    doc = Document()
    
    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # ── Title ──
    title_para = doc.add_heading('', level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('操作录屏报告')
    _set_run_font(run, font_size=Pt(26), color=RGBColor(0x1a, 0x3a, 0x5c))
    
    # ── Subtitle ──
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Operation Recording Report')
    _set_run_font(run, font_size=Pt(12), color=RGBColor(0x66, 0x66, 0x66))
    
    doc.add_paragraph()  # spacer
    
    # ── Info table ──
    info_table = doc.add_table(rows=3, cols=2, style='Light Shading')
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("项目名称", project_name or "N/A"),
        ("操作步骤", f"{len(groups)} 个用户动作"),
        ("截图数量", str(ss_count)),
    ]
    for i, (label, value) in enumerate(info_data):
        cell_l = info_table.cell(i, 0)
        cell_r = info_table.cell(i, 1)
        cell_l.text = label
        cell_r.text = value
        for cell in [cell_l, cell_r]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    
    doc.add_paragraph()  # spacer
    
    # ── Steps ──
    used_screenshots = set()
    report_base = os.path.dirname(output_path)
    
    for idx, g in enumerate(groups, 1):
        ts = g.get("timestamp", "")
        win = g.get("window", "")
        msg = g.get("message", "")
        desc = summarize_event(msg)
        coords = extract_coordinates(msg)
        ss = g.get("screenshot", "")
        
        # Step heading
        h = doc.add_heading(f'步骤 {idx}', level=2)
        run = h.runs[0]
        _set_run_font(run, color=RGBColor(0x1a, 0x3a, 0x5c))
        
        # Timestamp
        p = doc.add_paragraph()
        run = p.add_run(f'时间: {ts}')
        _set_run_font(run, font_size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))
        p.paragraph_format.space_after = Pt(4)
        
        # Detail table
        rows_data = [
            ("操作类型", desc),
            ("原始消息", msg),
            ("活动窗口", win),
        ]
        if coords:
            rows_data.append(("坐标", f"({coords[0]}, {coords[1]})"))
        
        step_table = doc.add_table(rows=len(rows_data), cols=2, style='Light Shading')
        step_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, value) in enumerate(rows_data):
            cell_l = step_table.cell(i, 0)
            cell_r = step_table.cell(i, 1)
            run_l = cell_l.paragraphs[0].add_run(label)
            _set_run_font(run_l, font_size=Pt(9), bold=True)
            run_r = cell_r.paragraphs[0].add_run(value)
            _set_run_font(run_r, font_size=Pt(9))
        
        # Screenshot
        if ss and ss not in used_screenshots:
            ss_full = os.path.join(report_base, ss)
            if os.path.exists(ss_full):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run()
                    run.add_picture(ss_full, width=Inches(6.0))
                    used_screenshots.add(ss)
                except Exception:
                    pass
        
        doc.add_paragraph()  # spacer between steps
    
    # ── Footer ──
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('报告由 AgentRunner Recorder 自动生成')
    _set_run_font(run, font_size=Pt(9), color=RGBColor(0x99, 0x99, 0x99))
    
    doc.save(output_path)
    return output_path


# ══════════════════════════════════════════════════════════════════════
# 统一入口
# ══════════════════════════════════════════════════════════════════════

def generate_reports(log_path: str, screenshot_dir: str, output_dir: str,
                     project_name: str = "", video_path: str = "") -> Dict[str, Optional[str]]:
    """Generate all three report formats (MD / HTML / DOCX).
    
    Returns dict with keys 'markdown', 'html', 'word' mapping to output paths.
    """
    events = parse_log(log_path)
    if not events:
        print("[report_generator] No events found in log")
        return {}
    
    result = {}
    
    # 1. Markdown
    md_path = os.path.join(output_dir, f"report_{project_name}.md")
    try:
        generate_markdown(events, screenshot_dir, md_path, project_name, video_path)
        result["markdown"] = md_path
        print(f"[report_generator] Markdown report: {md_path}")
    except Exception as e:
        print(f"[report_generator] Markdown error: {e}")
    
    # 2. HTML
    html_path = os.path.join(output_dir, f"report_{project_name}.html")
    try:
        generate_html(events, screenshot_dir, html_path, project_name, video_path)
        result["html"] = html_path
        print(f"[report_generator] HTML report: {html_path}")
    except Exception as e:
        print(f"[report_generator] HTML error: {e}")
    
    # 3. Word
    docx_path = os.path.join(output_dir, f"report_{project_name}.docx")
    try:
        generate_word(events, screenshot_dir, docx_path, project_name, video_path)
        result["word"] = docx_path
        print(f"[report_generator] Word report: {docx_path}")
    except Exception as e:
        print(f"[report_generator] Word error: {e}")
    
    # 4. JSON
    json_path = os.path.join(output_dir, f"report_{project_name}.json")
    try:
        generate_json(events, screenshot_dir, json_path, project_name, video_path)
        result["json"] = json_path
        print(f"[report_generator] JSON report: {json_path}")
    except Exception as e:
        print(f"[report_generator] JSON error: {e}")
    
    return result
